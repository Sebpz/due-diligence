# asset-mgmt-server.py
from fastmcp import FastMCP
import logging
from typing import List, Dict, Optional
from googlesearch import search
import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
import yfinance as yf

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

mcp = FastMCP("asset-mgmt-tools")


@mcp.tool()
def get_company_data(ticker_symbol: str) -> dict:
    """Fetch company info + last 5 days of stock prices."""
    if not isinstance(ticker_symbol, str) or not ticker_symbol.strip():
        return {"error": "Invalid ticker symbol", "info": {}, "historical_prices": {"index": [], "data": []}}

    try:
        company = yf.Ticker(ticker_symbol)
        company_info = company.info or {}
        if not company_info:
            return {"error": "No company information found", "info": {}, "historical_prices": {"index": [], "data": []}}

        historical_data = company.history(period="5d")
        historical_dict = {
            'index': historical_data.index.strftime('%Y-%m-%d').tolist(),
            'data': historical_data.to_dict('records')
        } if not historical_data.empty else {"index": [], "data": []}

        return {"info": company_info, "historical_prices": historical_dict}

    except Exception as e:
        logger.error(f"Error retrieving data for {ticker_symbol}: {e}")
        return {"error": str(e), "info": {}, "historical_prices": {"index": [], "data": []}}

@mcp.tool()
def create_professional_report(filename: str, report_data: dict, table_data: Optional[dict] = None) -> dict:
    """Generate a professional Word doc with sections + optional table."""
    try:
        if not filename.endswith('.docx'):
            raise ValueError("Filename must end with .docx")

        document = Document()
        document.add_heading(report_data.get("title", "Generated Report"), level=0)

        for section in report_data.get("sections", []):
            if isinstance(section, dict):
                document.add_heading(section.get("heading", ""), level=1)
                document.add_paragraph(section.get("body_text", ""))

        if table_data and isinstance(table_data, dict):
            document.add_heading("Financial Data", level=2)
            columns = table_data.get('columns', [])
            data = table_data.get('data', [])
            if columns and data:
                table = document.add_table(rows=1, cols=len(columns))
                table.style = 'Table Grid'
                hdr_cells = table.rows[0].cells
                for i, col in enumerate(columns):
                    hdr_cells[i].text = str(col)
                for row_data in data:
                    row_cells = table.add_row().cells
                    for i, col in enumerate(columns):
                        row_cells[i].text = str(row_data.get(col, ''))

        if "footer" in report_data:
            section = document.sections[0]
            paragraph = section.footer.paragraphs[0]
            paragraph.text = report_data["footer"]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        document.save(filename)
        return {"status": "success", "filename": filename}

    except Exception as e:
        return {"status": "error", "message": str(e)}

@mcp.tool()
def search_for_news(query: str) -> List[Dict[str, str]]:
    """Search Google for recent news URLs."""
    if not query.strip():
        return []
    try:
        search_results = list(search(f"{query} news", num_results=10))
        return [{"headline": url, "url": url} for url in search_results if isinstance(url, str)]
    except Exception as e:
        logger.error(f"Error searching for news: {e}")
        return []

@mcp.tool()
def get_soup_from_url(url: str) -> Dict[str, str]:
    """Scrape a webpage's title, text, and HTML."""
    if not url.startswith(('http://', 'https://')):
        return {"title": "", "text": "", "html": "", "error": "Invalid URL format"}

    try:
        r = requests.get(url, timeout=10, headers={'User-Agent': 'Mozilla/5.0'})
        r.raise_for_status()
        soup = BeautifulSoup(r.text, 'html.parser')
        text_content = ' '.join(soup.get_text(separator=' ', strip=True).split())
        return {
            "title": soup.title.string if soup.title else "",
            "text": text_content,
            "html": str(soup),
            "status": "success"
        }
    except Exception as e:
        return {"title": "", "text": "", "html": "", "error": str(e), "status": "error"}

if __name__ == "__main__":
    mcp.run(transport="http", host="0.0.0.0", port=8000)
